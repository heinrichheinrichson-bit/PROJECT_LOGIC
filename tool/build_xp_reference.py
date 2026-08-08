from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Project Logic XP-Testreferenz.docx"

INK = RGBColor(25, 36, 35)
TEAL = RGBColor(0, 121, 107)
BLUE = RGBColor(46, 91, 167)
MUTED = RGBColor(82, 96, 94)
LIGHT = "E8EEF5"
TEAL_LIGHT = "DDEFEA"
WHITE = RGBColor(255, 255, 255)


def set_font(run, size=11, bold=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def style_table(table, header=True):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_index == 0:
                shade(cell, LIGHT)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.15
                for run in paragraph.runs:
                    set_font(run, size=9.5, bold=header and row_index == 0)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_note(doc, label, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    shade(table.cell(0, 0), TEAL_LIGHT)
    p = table.cell(0, 0).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(f"{label}: ")
    set_font(r, bold=True, color=TEAL)
    r = p.add_run(text)
    set_font(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, before, after, color in (
        ("Heading 1", 16, 18, 10, BLUE),
        ("Heading 2", 13, 14, 7, BLUE),
        ("Heading 3", 12, 10, 5, RGBColor(31, 77, 120)),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header.add_run("PROJECT LOGIC  ·  TESTREFERENZ"), size=8.5, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(footer.add_run("XP-Balance v3 · Stand 8. August 2026"), size=8.5, color=MUTED)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(2)
    set_font(kicker.add_run("TESTREFERENZ"), size=9, bold=True, color=TEAL)
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    set_font(title.add_run("Project Logic – XP-Werte"), size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    set_font(subtitle.add_run("Was beim Testen wie viele Erfahrungspunkte bringt"), size=13, color=MUTED)

    add_note(doc, "Aktiv", "Diese Werte sind im aktuellen Build zentral hinterlegt. Bereits früher gespeicherte XP werden nicht nachträglich verringert oder überschrieben.")

    add_heading(doc, "1. Rätselabschlüsse", 1)
    doc.add_paragraph("Die Tabelle zeigt die XP für einen ersten Abschluss mit mindestens einem verwendeten Hinweis. Ein Abschluss ohne Hinweis erhält zusätzlich 10 XP.")
    rows = [
        ("Spielart", "Leicht", "Mittel", "Schwer"),
        ("Rätselsammlung", "20", "30", "45"),
        ("Zufallsrätsel", "25", "35", "50"),
        ("Tagesrätsel / Kalender", "35", "45", "60"),
        ("Tutorial", "10", "20", "35"),
    ]
    table = doc.add_table(rows=len(rows), cols=4)
    for r, values in enumerate(rows):
        for c, value in enumerate(values):
            table.cell(r, c).text = value
    set_table_geometry(table, [3960, 1800, 1800, 1800])
    style_table(table)

    add_heading(doc, "Bonus- und Wiederholungsregeln", 2)
    bullets = [
        "Ohne Hinweis gelöst: +10 XP.",
        "Erneut gelöstes Sammlungs- oder Zufallsrätsel: 5 XP pauschal.",
        "Erneut gelöstes Tagesrätsel: 0 XP. Der Kalenderstatus bleibt natürlich erhalten.",
        "Werbehinweise, Testdialoge und das bloße Öffnen eines Rätsels geben keine XP.",
        "Rastergröße gibt derzeit keinen eigenen XP-Aufschlag; die Schwierigkeit bildet den Aufwand ab.",
    ]
    for text in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    add_heading(doc, "Schnelle Testbeispiele", 2)
    examples = [
        ("Leichtes Sammlungsrätsel ohne Hinweis", "20 + 10", "30 XP"),
        ("Mittleres Zufallsrätsel mit Hinweis", "35 + 0", "35 XP"),
        ("Schweres Tagesrätsel ohne Hinweis", "60 + 10", "70 XP"),
        ("Dasselbe Tagesrätsel erneut", "Wiederholung", "0 XP"),
    ]
    table = doc.add_table(rows=1, cols=3)
    table.rows[0].cells[0].text = "Testfall"
    table.rows[0].cells[1].text = "Rechnung"
    table.rows[0].cells[2].text = "Ergebnis"
    for values in examples:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    set_table_geometry(table, [5400, 1800, 2160])
    style_table(table)

    doc.add_page_break()
    add_heading(doc, "2. Erfolge", 1)
    doc.add_paragraph("Erfolgs-XP werden genau einmal vergeben. Schon gutgeschriebene Werte bleiben als unveränderliches Ereignis im XP-Konto gespeichert.")
    achievements = [
        ("Erster Abschluss", "Der Anfang ist gemacht", 25),
        ("10 / 50 / 100 Rätsel", "Meilensteine", "50 / 100 / 175"),
        ("250 / 500 / 1.000 Rätsel", "Langzeit-Meilensteine", "250 / 400 / 750"),
        ("Erstes Rätsel je Spielart", "6 einzelne Erfolge", "je 25"),
        ("3 / 7 / 30 Tage Serie", "Streak-Erfolge", "40 / 90 / 250"),
        ("7 / 30 Tagesrätsel", "Kalender-Erfolge", "75 / 200"),
        ("10 / 50 Zufallsrätsel", "Generator-Erfolge", "75 / 200"),
        ("Alle Spielarten entdeckt", "Vielseitiger Denker", 150),
        ("1 / 10 Stunden Spielzeit", "Zeit-Erfolge", "75 / 250"),
        ("Gesamte Sammlung gelöst", "Alles gesehen", 400),
        ("5 schwere Rätsel", "Harte Nüsse", 100),
        ("Raster ab 10 × 10", "Großes Raster", 50),
    ]
    table = doc.add_table(rows=1, cols=3)
    for index, value in enumerate(("Bedingung", "Erfolg", "XP")):
        table.cell(0, index).text = value
    for condition, title_text, points in achievements:
        cells = table.add_row().cells
        cells[0].text = str(condition)
        cells[1].text = str(title_text)
        cells[2].text = str(points)
    set_table_geometry(table, [3960, 3600, 1800])
    style_table(table)

    add_heading(doc, "3. Tages- und Wochenziele", 1)
    doc.add_paragraph("Die sichtbaren Ziele wechseln reproduzierbar mit dem Datum beziehungsweise der Kalenderwoche. Ein Ziel kann wegen seiner eindeutigen Kennung nur einmal XP vergeben.")
    mission_rows = [
        ("Ein einzelnes Tagesziel", "5 XP"),
        ("Alle drei Tagesziele", "+15 XP Komplettbonus"),
        ("Ein einzelnes Wochenziel", "40 XP"),
        ("Beide Wochenziele", "+30 XP Komplettbonus"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Abschluss"
    table.rows[0].cells[1].text = "Belohnung"
    for label, points in mission_rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = points
    set_table_geometry(table, [6840, 2520])
    style_table(table)

    add_heading(doc, "Rotation und Testverhalten", 2)
    mission_notes = [
        "Heute: Das Tagesrätsel ist immer dabei; zwei weitere Ziele rotieren zwischen Zufallsrätsel, Sammlung, ohne Hinweis und zwei Spielarten.",
        "Diese Woche: Zwei Ziele rotieren zwischen drei aktiven Tagen, drei Spielarten, fünf Rätseln, drei Lösungen ohne Hinweis und zwei schweren Rätseln.",
        "Alle Fortschritte werden aus dem Abschlussverlauf der jeweiligen Tages- oder Wochenperiode berechnet; Lebenszeitwerte zählen nicht versehentlich in Wochenziele hinein.",
        "Im XP-Verlauf erscheinen Einzelziele als Missionsbelohnung und vollständige Zielsets als eigener Komplettbonus.",
    ]
    for text in mission_notes:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    add_heading(doc, "4. Noch geplant – derzeit 0 XP", 1)
    planned = [
        "Streak-Eis, Erinnerungen und rückwirkendes Einfrieren: keine XP-Kaufmöglichkeit; Motivation statt Pay-to-win.",
        "Saisonale Ereignisse: später separate Event-XP, erst wenn das Ereignissystem existiert.",
        "Premiumkauf und Werbung: niemals direkte XP. Premium entfernt Reibung, kauft aber keinen Fortschritt.",
    ]
    for text in planned:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.add_run(text)

    add_note(doc, "Testhinweis", "Nach einem Abschluss sollte sich die Levelanzeige aktualisieren. Bei einem identischen Tagesrätsel darf die zweite Lösung keine weiteren XP geben. Sammlungs- und Zufallswiederholungen geben nur 5 XP.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
